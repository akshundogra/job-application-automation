import sys
import json
import os
import shutil
import subprocess
import re
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_LINE_SPACING

def add_hyperlink(paragraph, url, text, bold=False):
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    if bold:
        b = OxmlElement('w:b')
        rPr.append(b)
    new_run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def fix_utm_in_table(doc, utm_url):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    full_text = para.text
                    if 'Your_Domain' in full_text:
                        p_xml = para._p
                        hyperlinks = p_xml.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hyperlink')
                        for hl in hyperlinks:
                            hl_text = ''.join(t.text for t in hl.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t') if t.text)
                            if 'Your_Domain' in hl_text and 'utm' not in hl_text:
                                r_id = hl.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                                if r_id:
                                    try:
                                        para.part.rels[r_id]._target = utm_url
                                        for t_elem in hl.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
                                            t_elem.text = 'Your_Domain'
                                            break
                                    except:
                                        pass


def apply_changes_to_resume(source_path, output_path, changes, utm_url):
    shutil.copy2(source_path, output_path)
    doc = Document(output_path)

    for change in changes:
        find_text = change.get('find')
        replace_text = change.get('replace', '')

        if find_text is None or str(find_text).lower() == 'null' or str(find_text).strip() == '':
            insert_skills_line(doc, replace_text)
            continue

        for para in doc.paragraphs:
            if find_text in para.text:
                for run in para.runs:
                    run.text = ''
                if para.runs:
                    para.runs[0].text = replace_text
                else:
                    para.add_run(replace_text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if find_text in para.text:
                            for run in para.runs:
                                run.text = ''
                            if para.runs:
                                para.runs[0].text = replace_text
                            else:
                                para.add_run(replace_text)

    fix_utm_in_table(doc, utm_url)
    doc.save(output_path)
    print(f"Resume saved: {output_path}")

def add_formatted_run(paragraph, text):
    parts = re.split(r'\*\*(.+?)\*\*', text)
    for i, part in enumerate(parts):
        if not part:
            continue
        run = paragraph.add_run(part)
        run.bold = (i % 2 == 1)
        run.font.size = Pt(10.5)
        run.font.name = 'Calibri'

def create_cover_letter_docx(cover_text, output_path, utm_url):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)

    lines = cover_text.strip().split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            continue
        if line.startswith('•') or line.startswith('-'):
            content = line.lstrip('•- ').strip()
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_run(p, content)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            p.paragraph_format.line_spacing = 1.1
            continue
        if 'YOUR_Domain' in line:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            if ':' in line and 'YOUR_Domain' in line:
                prefix = line.split('YOUR_Domain')[0]
                run = p.add_run(prefix)
                run.bold = True
                run.font.size = Pt(10.5)
                run.font.name = 'Calibri'
                add_hyperlink(p, utm_url, 'YOUR_Domain', bold=True)
            else:
                add_hyperlink(p, utm_url, 'YOUR_Domain', bold=True)
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.1
        add_formatted_run(p, line)

    doc.save(output_path)
    print(f"Cover letter saved: {output_path}")

def convert_to_pdf_with_word(docx_path):
    pdf_path = docx_path.replace('.docx', '.pdf')
    output_dir = os.path.dirname(docx_path)
    subprocess.run([
        '/opt/homebrew/bin/soffice',
        '--headless',
        '--convert-to', 'pdf',
        '--outdir', output_dir,
        docx_path
    ], check=True)
    print(f"PDF created: {pdf_path}")
    return pdf_path

def scrub_metadata(pdf_path, author_name):
    subprocess.run([
        '/opt/homebrew/bin/exiftool',
        f'-Author={author_name}',
        '-Creator=Microsoft Word',
        '-Producer=Microsoft Word for Microsoft 365',
        '-Title=',
        '-Subject=',
        '-Keywords=',
        '-overwrite_original',
        pdf_path
    ], check=True)
    print(f"Metadata scrubbed: {pdf_path}")

def main():
    config_path = sys.argv[1]
    with open(config_path, 'r', encoding='utf-8') as f:
        config = json.load(f)

    language = config['language']
    changes = config['changes']
    cover_text = config['coverText']
    file_base = config['fileBase']
    output_dir = config['outputDir']
    utm_resume = config['utmResume']
    utm_cover = config['utmCover']

    date_prefix = datetime.now().strftime('%Y%m%d')
    os.makedirs(output_dir, exist_ok=True)

    if 'English' in language:
        source_resume = '/Users/YOUR_USERNAME/.n8n-files/YOUR_RESUME-EN.docx'
        resume_out = os.path.join(output_dir, f'{date_prefix}_{file_base}_Resume_EN.docx')
        cover_out = os.path.join(output_dir, f'{date_prefix}_{file_base}_CoverLetter_EN.docx')
    else:
        source_resume = '/Users/YOUR_USERNAME/.n8n-files/YOUR_RESUME-DE'
        resume_out = os.path.join(output_dir, f'{date_prefix}_{file_base}_Lebenslauf_DE.docx')
        cover_out = os.path.join(output_dir, f'{date_prefix}_{file_base}_Anschreiben_DE.docx')

    apply_changes_to_resume(source_resume, resume_out, changes, utm_resume)
    create_cover_letter_docx(cover_text, cover_out, utm_cover)
    resume_pdf = convert_to_pdf_with_word(resume_out)
    cover_pdf = convert_to_pdf_with_word(cover_out)
    scrub_metadata(resume_pdf, 'Your_Name')
    scrub_metadata(cover_pdf, 'Your_Name')

    result = {
        'resumePDF': resume_pdf,
        'coverPDF': cover_pdf,
        'resumeDOCX': resume_out,
        'coverDOCX': cover_out
    }
    print(json.dumps(result))

if __name__ == '__main__':
    main()
